import tkinter as tk
from tkinter import filedialog
from tkinter import ttk
import csv
import datetime
import docx

window = tk.Tk()
window.geometry('700x300')
window.title('Convertidor CSV a word')
day = datetime.datetime.now().strftime('%Y-%m-%d')

def convert(filepath):

    with open(filepath, encoding='utf-8') as file:
        reader = csv.reader(file)
        respuestas = list(reader)

    doc = docx.Document()

    for i in range(1,len(respuestas)):
        for j in range(1,len(respuestas[i])):
            doc.add_heading(respuestas[0][j], 4)
            paragraph = doc.add_paragraph(respuestas[i][j])
        run = paragraph.add_run()
        run.add_break(docx.enum.text.WD_BREAK.PAGE)

    doc.save(salida)

    etiqueta = tk.Label(window, text = f'Archivo generado: {salida}')  # Declaro la etiqueta
    etiqueta.grid(row = 3, column = 1)  # Agrega etiqueta a la parte de abajo de la ventana

def archivo():
    window.filename =  filedialog.askopenfilename()
    print(window.filename)
    global ruta
    ruta = window.filename
    global salida
    salida = '.'.join(ruta.split('.')[:-1]) + f' {day}.docx'

    etiquetaArchivo = tk.Label(window, text = ruta)  # Declaro la etiqueta
    etiquetaArchivo.grid(row = 2, column = 1)  # Agrega etiqueta a la parte de abajo de la ventana

    # Botón para correr el programa
    excecute_button = tk.Button(window, text = 'Ejecutar', command = lambda: convert(ruta))
    excecute_button.grid(row = 3, column = 0)

# Botón para seleccionar archivo
file_button = tk.Button(window, text = 'Seleccione archivo', command = archivo)
file_button.grid(row = 2, column = 0)

nombre = tk.Label(window, text = r'Matias Monticelli')
nombre.grid(row = 4, column = 1)

linkedin = tk.Label(window, text = r'https://www.linkedin.com/in/matiasmonticelli/')
linkedin.grid(row = 5, column = 1)

window.mainloop()